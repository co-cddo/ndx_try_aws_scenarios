"""Generate fake parish-flavoured sample documents for the Paperless-ngx demo.

Runs once in the init container on first boot. Writes ~36 documents into
/efs/consume so Paperless picks them up automatically and feeds the AI
classification pipeline a rich variety of inputs.

Mix:
  - 12 planning notices (full / listed building / TPO / discharge)
  - 8 invoices (grounds, water, insurance, audit, IT, hall hire, etc.)
  - 6 council meeting minutes (monthly, Oct 2025 - Mar 2026)
  - 3 agendas (Apr / May / Jun 2026)
  - 5 correspondence emails (.eml)
  - 2 statutory reports (annual return, asset register)

Uses reportlab for PDFs and python-docx for .docx. .eml files are plain text.
The OCR + downstream Bedrock features then run on each one as they're consumed.
"""

import os
import sys
import textwrap
from datetime import date

OUT_DIR = "/efs/consume"
os.makedirs(OUT_DIR, exist_ok=True)


def write_pdf(filename: str, title: str, body: str) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

    doc = SimpleDocTemplate(os.path.join(OUT_DIR, filename), pagesize=A4)
    styles = getSampleStyleSheet()
    story = [Paragraph(title, styles["Title"]), Spacer(1, 12)]
    for para in body.strip().split("\n\n"):
        story.append(Paragraph(para.replace("\n", "<br/>"), styles["BodyText"]))
        story.append(Spacer(1, 8))
    doc.build(story)


def write_docx(filename: str, title: str, body: str) -> None:
    from docx import Document

    d = Document()
    d.add_heading(title, level=0)
    for para in body.strip().split("\n\n"):
        d.add_paragraph(para)
    d.save(os.path.join(OUT_DIR, filename))


def write_eml(filename: str, subject: str, sender: str, body: str, when: str = None) -> None:
    when = when or date.today().isoformat()
    eml = textwrap.dedent(f"""\
        From: {sender}
        To: clerk@meadowbrook-pc.gov.uk
        Subject: {subject}
        Date: {when}
        Content-Type: text/plain; charset=utf-8

        {body.strip()}
        """)
    with open(os.path.join(OUT_DIR, filename), "w") as f:
        f.write(eml)


# ---------- Planning notices (12) ----------

PLANNING_TEMPLATE = """Notice is hereby given that an application has been submitted to the
Local Planning Authority by {applicant} in respect of:

Application Number: {ref}
Site Address: {address}
Description: {description}

{type_specific}

Members of the public may inspect the application and submit representations to
the Planning Department within 21 days of the date of this notice. Comments may
be submitted online via the council's planning portal or in writing to the
Planning Department at the address above.

Date of notice: {issued}
Issued by: Meadowbrook Parish Council on behalf of the Local Planning Authority"""


def planning_notice(ref, address, description, applicant, issued, type_specific=""):
    safe = ref.replace("/", "-")
    title = f"Town and Country Planning Act 1990 - Planning Notice {ref}"
    body = PLANNING_TEMPLATE.format(
        ref=ref,
        address=address,
        description=description,
        applicant=applicant,
        issued=issued,
        type_specific=type_specific,
    )
    write_pdf(f"planning_notice_{safe}.pdf", title, body)


def write_planning_notices():
    planning_notice(
        "24/01234/FUL", "The Old Forge, Mill Lane, Meadowbrook, MB4 8XR",
        "Erection of a single-storey rear extension following demolition of existing conservatory.",
        "Mr R. Hawthorne", "12 March 2026",
    )
    planning_notice(
        "24/01987/LBC", "7 Church Walk, Meadowbrook, MB4 7QP (Grade II listed)",
        "Replacement of rotten timber sash windows on principal elevation with like-for-like single-glazed sash windows in painted hardwood.",
        "Ms E. Whitaker", "18 March 2026",
        "This application is for Listed Building Consent. The statutory consultation period is 21 days from 18 March 2026. Comments must reference 24/01987/LBC and be submitted to the Conservation Officer at conservation@meadowbrook.gov.uk.",
    )
    planning_notice(
        "24/02014/TPO", "Land at the rear of 14 Beech Drive, Meadowbrook, MB4 8AA",
        "Application to fell one mature horse chestnut covered by Tree Preservation Order MB-TPO-1998-024 on grounds of subsidence damage to neighbouring property.",
        "Mrs S. Patel", "22 March 2026",
        "An arboricultural report has been submitted in support. The TPO consultation runs for 28 days.",
    )
    planning_notice(
        "24/02055/FUL", "Bramble Cottage, Heath Road, Meadowbrook, MB4 9JL",
        "Change of use of redundant agricultural barn to a single dwelling (Class C3) including conversion works and associated parking.",
        "Mr & Mrs J. Adair", "26 March 2026",
    )
    planning_notice(
        "24/02091/HOU", "12 Willow Crescent, Meadowbrook, MB4 8DG",
        "Erection of a detached double garage with home office above.",
        "Dr M. Okonkwo", "29 March 2026",
    )
    planning_notice(
        "24/02112/CLD", "Meadowbrook Service Station, A1198 Trunk Road, Meadowbrook, MB4 6PT",
        "Certificate of lawfulness for proposed use of forecourt for sale of food and hot drinks from existing kiosk between 06:00 and 22:00 daily.",
        "Highstone Energy Retail Ltd", "02 April 2026",
    )
    planning_notice(
        "24/02134/ADV", "Meadowbrook Village Hall, Green Lane, Meadowbrook, MB4 7FS",
        "Display of one externally illuminated post-mounted noticeboard at the entrance to the village hall car park.",
        "Meadowbrook Village Hall Trustees", "05 April 2026",
    )
    planning_notice(
        "24/02158/LBC", "St Mary's Church, Church Walk, Meadowbrook, MB4 7QQ (Grade I listed)",
        "Internal re-ordering of nave to remove fixed pews and install reversible chairs to enable wider community use.",
        "PCC of St Mary's Meadowbrook", "08 April 2026",
        "Listed Building Consent. Application accompanied by a Statement of Significance and a heritage impact assessment.",
    )
    planning_notice(
        "24/02181/DOC", "Meadowbrook Primary School, School Lane, Meadowbrook, MB4 8RR",
        "Application for discharge of conditions 4 (materials), 7 (landscaping) and 9 (drainage) of permission 23/01882/FUL.",
        "Cambridgeshire County Council", "11 April 2026",
    )
    planning_notice(
        "24/02199/HOU", "Yew Tree House, 3 The Green, Meadowbrook, MB4 7DA",
        "Replacement of existing rear conservatory with a single-storey orangery extension.",
        "Mr P. Nightingale", "14 April 2026",
    )
    planning_notice(
        "24/02208/AGR", "Meadowbrook Farm, Fen Drove, Meadowbrook, MB4 9QR",
        "Erection of a steel-framed agricultural building for storage of fodder and machinery (prior approval).",
        "Meadowbrook Farms Ltd", "17 April 2026",
    )
    planning_notice(
        "24/02221/FUL", "The Old Bakery, 22 High Street, Meadowbrook, MB4 7AB",
        "Change of use from retail (E(a)) to mixed retail and cafe (E(a)/(b)) with associated minor external alterations.",
        "Ms L. Beckett", "20 April 2026",
    )


# ---------- Invoices (8) ----------

def invoice(filename, title, body):
    write_pdf(filename, title, body)


def write_invoices():
    invoice(
        "invoice_dunbar_grounds_2026-03.pdf",
        "Invoice DG-2026-0341",
        """DUNBAR GROUNDS MAINTENANCE LIMITED
Unit 4, Brookside Industrial Estate, Meadowbrook, MB4 6TT
VAT no. GB 412 8856 21

Invoice to: Meadowbrook Parish Council
Invoice number: DG-2026-0341
Date: 27 March 2026
Payment terms: 30 days

Description                                              Amount (GBP)
Grass cutting - village green and recreation field           320.00
Hedge trimming - cemetery boundary                            85.00
Litter pick - High Street and bus stops (March visits)        62.50

Subtotal                                                     467.50
VAT @ 20%                                                     93.50
TOTAL DUE                                                    561.00

Please pay by BACS to Lloyds 30-99-12 account 1234 5678,
quoting reference DG-2026-0341.""",
    )
    invoice(
        "invoice_dunbar_grounds_2026-04.pdf",
        "Invoice DG-2026-0412",
        """DUNBAR GROUNDS MAINTENANCE LIMITED
Unit 4, Brookside Industrial Estate, Meadowbrook, MB4 6TT
VAT no. GB 412 8856 21

Invoice to: Meadowbrook Parish Council
Invoice number: DG-2026-0412
Date: 24 April 2026
Payment terms: 30 days

Description                                              Amount (GBP)
Grass cutting - village green and recreation field (x2)      640.00
Spring border weeding - Memorial Garden                      120.00
Tree stake replacement - new whips on cemetery boundary       45.00

Subtotal                                                     805.00
VAT @ 20%                                                    161.00
TOTAL DUE                                                    966.00

Direct debit collection on 24 May 2026.""",
    )
    invoice(
        "invoice_anglia_water_2026-03.pdf",
        "Anglia Water Services - Invoice 26/2255198",
        """Anglia Water Services Limited
PO Box 91, Cambridge CB1 9YY

Account holder: Meadowbrook Parish Council (Pavilion supply)
Account number: 8856-2241-9
Invoice number: 26/2255198
Period: 01 January 2026 - 31 March 2026

Water used: 14 cubic metres
Wastewater: 14 cubic metres
Standing charge:                                              28.40
Volume charge (water + wastewater):                           42.84
Total before VAT                                              71.24
VAT (water services - exempt)                                  0.00
TOTAL DUE                                                     71.24

Direct debit collection on 18 April 2026.""",
    )
    invoice(
        "invoice_zurich_insurance_2026-04.pdf",
        "Zurich Municipal - Annual Premium Renewal",
        """Zurich Municipal
Renewal Notice 2026/27 - Combined Liability and Property
Policy Reference: ZM/PC/884412/26

Policyholder: Meadowbrook Parish Council
Renewal date: 01 May 2026
Cover period: 01 May 2026 to 30 April 2027

Section A - Public Liability (£10m limit):                  482.00
Section B - Employers Liability (£10m limit):                94.50
Section C - Property (Pavilion + village hall content):      275.00
Section D - Money & legal expenses:                           48.00
Section E - Officials indemnity:                             125.00

Net premium                                                1,024.50
Insurance Premium Tax @ 12%                                  122.94
TOTAL PAYABLE                                              1,147.44

Please confirm renewal by 24 April 2026.""",
    )
    invoice(
        "invoice_internal_audit_2026.pdf",
        "Goodwin & Hayes - Internal Audit Fee 2025/26",
        """Goodwin & Hayes Audit Services
14 Eastgate, Cambridge, CB2 1JA
Registered with NALC as a competent internal auditor

Invoice to: Meadowbrook Parish Council
Invoice number: GH/2026/441
Date: 30 April 2026

For: Annual internal audit of accounting records, governance arrangements
and effectiveness of internal control for year ended 31 March 2026 in line
with the Practitioners' Guide for Local Councils.

Audit fee:                                                   480.00
Travel and disbursements:                                     35.00
TOTAL DUE                                                    515.00

Payment terms: 30 days. Bank details on remittance advice attached.""",
    )
    invoice(
        "invoice_drupal_hosting_2026-q1.pdf",
        "Cloudtide Ltd - Website Hosting Q1 2026",
        """Cloudtide Ltd
Office 12, The Old Mill, Burwell, Cambs, CB25 0AB

Invoice to: Meadowbrook Parish Council
Invoice number: CT-MB-2026-Q1
Quarter: January - March 2026

Hosting (managed Drupal, 4 GB plan):                         180.00
SSL certificate renewal:                                       0.00
Domain registration meadowbrook-pc.gov.uk:                    24.00
Email forwarding (5 mailboxes):                               36.00
Subtotal                                                     240.00
VAT @ 20%                                                     48.00
TOTAL                                                        288.00

Payment terms: 14 days from date of invoice.""",
    )
    invoice(
        "invoice_villagehall_hire_2026-03.pdf",
        "Village Hall Trustees - Hire of Hall (March)",
        """Meadowbrook Village Hall Trustees
Registered Charity 1108412

Invoice to: Meadowbrook Parish Council
Invoice number: VH-2026-014
Date: 31 March 2026

For: Hire of main hall for Full Council meetings 11 March 2026
and Planning Committee 18 March 2026.

3 hours x 11 March (main hall):                               33.00
2 hours x 18 March (small room):                              16.00
TOTAL DUE                                                     49.00

Trustees thank the Council for its continued support.""",
    )
    invoice(
        "invoice_acorn_skip_hire_2026-04.pdf",
        "Acorn Skip Hire - Cemetery clearance",
        """Acorn Skip Hire (East Anglia) Ltd
Yard 3, Long Drove, March, PE15 0LZ

Invoice to: Meadowbrook Parish Council
Invoice number: AC/8821
Date: 14 April 2026

Hire of one 6-yard skip, delivered 09 April, collected 14 April,
for clearance of dead vegetation and old kerb stones from cemetery
extension area.

Skip hire (5 working days):                                  185.00
Disposal of mixed inert waste (1.2 tonnes):                   78.00
Subtotal                                                     263.00
VAT @ 20%                                                     52.60
TOTAL                                                        315.60

Please pay by 14 May 2026.""",
    )


# ---------- Council minutes (6) ----------

def write_minutes():
    write_docx(
        "minutes_full_council_2025-10-08.docx",
        "Meadowbrook Parish Council - Minutes of Full Council, 8 October 2025",
        """Present: Cllr S. Whitfield (Chair), Cllr A. Bramley, Cllr P. Okafor,
Cllr H. Mendez, Cllr R. Goodwin. In attendance: Mrs C. Wright (Clerk).
Apologies: Cllr E. Dawes, Cllr J. Patel.

25/10/01 Declarations of interest. None.

25/10/02 Minutes of 10 September 2025 approved as a true record.

25/10/03 Public participation. Two residents raised concerns about HGV
movements on Mill Lane during the autumn beet harvest. The Clerk will
write to the County Highways Manager.

25/10/04 Planning. The Council resolved to support 24/01821/HOU and to
make no comment on 24/01884/CLD.

25/10/05 Finance. Bank balance at 30 September: GBP 19,210.05. Payments
of GBP 612.40 (grounds) and GBP 71.24 (water) approved.

25/10/06 Cemetery extension - the working group reported progress on
boundary clearance. A skip will be hired in spring.

25/10/07 Date of next meeting: Wednesday 12 November 2025, 19:30.""",
    )
    write_docx(
        "minutes_full_council_2025-11-12.docx",
        "Meadowbrook Parish Council - Minutes of Full Council, 12 November 2025",
        """Present: Cllr S. Whitfield (Chair), Cllr A. Bramley, Cllr P. Okafor,
Cllr H. Mendez, Cllr E. Dawes, Cllr J. Patel. In attendance: Mrs C. Wright (Clerk).
Apologies: Cllr R. Goodwin.

25/11/01 Declarations of interest. Cllr Patel declared a non-pecuniary
interest in agenda item 6 (Village Hall hire fees).

25/11/02 Minutes of 8 October approved.

25/11/03 Public participation. None.

25/11/04 Planning. The Council noted the grant of permission for
24/01821/HOU and resolved to object to 24/01934/FUL on amenity grounds.

25/11/05 Remembrance Sunday. The Clerk thanked councillors and the
Royal British Legion for organising the wreath-laying ceremony, attended
by approximately 230 residents.

25/11/06 Village Hall hire fees - resolved to maintain current rates
into the 2026/27 financial year.

25/11/07 Precept for 2026/27. The Clerk presented draft figures showing
a band D equivalent of GBP 38.40, an increase of 1.6 per cent. Decision
deferred to December meeting following further review.

25/11/08 Date of next meeting: Wednesday 10 December 2025, 19:30.""",
    )
    write_docx(
        "minutes_full_council_2025-12-10.docx",
        "Meadowbrook Parish Council - Minutes of Full Council, 10 December 2025",
        """Present: Cllr S. Whitfield (Chair), Cllr A. Bramley, Cllr P. Okafor,
Cllr H. Mendez, Cllr R. Goodwin, Cllr E. Dawes. In attendance: Mrs C. Wright (Clerk).
Apologies: Cllr J. Patel.

25/12/01 Declarations of interest. None.

25/12/02 Minutes of 12 November approved.

25/12/03 Public participation. None.

25/12/04 Planning. No new applications.

25/12/05 Precept 2026/27. After discussion, the Council resolved to
set the precept at GBP 41,800 (band D equivalent GBP 38.40), an
increase of 1.6 per cent on 2025/26.

25/12/06 Christmas tree lighting - the Clerk reported a successful event
attended by approximately 180 residents and thanks were expressed to
the Cubs and Brownies for the carol singing.

25/12/07 Date of next meeting: Wednesday 14 January 2026, 19:30.""",
    )
    write_docx(
        "minutes_full_council_2026-01-14.docx",
        "Meadowbrook Parish Council - Minutes of Full Council, 14 January 2026",
        """Present: Cllr S. Whitfield (Chair), Cllr A. Bramley, Cllr P. Okafor,
Cllr H. Mendez, Cllr E. Dawes, Cllr J. Patel. In attendance: Mrs C. Wright (Clerk).
Apologies: Cllr R. Goodwin.

26/01/01 Declarations of interest. None.

26/01/02 Minutes of 10 December approved.

26/01/03 Public participation. A resident raised concerns about parking
on grass verges along School Lane. Referred to Highways.

26/01/04 Planning. The Council resolved to support 24/02014/HOU.

26/01/05 Recreation ground play equipment - quotations received from
two suppliers. Decision deferred pending a third.

26/01/06 Annual Parish Meeting - confirmed for 18 March 2026 at the
village hall.

26/01/07 Date of next meeting: Wednesday 11 February 2026, 19:30.""",
    )
    write_docx(
        "minutes_full_council_2026-02-11.docx",
        "Meadowbrook Parish Council - Minutes of Full Council, 11 February 2026",
        """Present: Cllr S. Whitfield (Chair), Cllr A. Bramley, Cllr P. Okafor,
Cllr H. Mendez, Cllr R. Goodwin, Cllr E. Dawes, Cllr J. Patel.
In attendance: Mrs C. Wright (Clerk).

26/02/01 Declarations of interest. None.

26/02/02 Minutes of 14 January approved.

26/02/03 Public participation. A representative of the Cricket Club
requested informal consent to install a new sightscreen on the
recreation ground. The Council asked for written specifications.

26/02/04 Planning. Resolved to make no comment on three minor householder
applications.

26/02/05 Recreation ground play equipment - third quotation received.
Decision deferred to next meeting.

26/02/06 Internal auditor appointment for 2025/26 confirmed: Goodwin &
Hayes Audit Services.

26/02/07 Date of next meeting: Wednesday 11 March 2026, 19:30.""",
    )
    write_docx(
        "minutes_full_council_2026-03-11.docx",
        "Meadowbrook Parish Council - Minutes of Full Council, 11 March 2026",
        """Present: Cllr S. Whitfield (Chair), Cllr A. Bramley, Cllr P. Okafor,
Cllr H. Mendez, Cllr J. Patel. In attendance: Mrs C. Wright (Clerk).
Apologies: Cllr R. Goodwin, Cllr E. Dawes.

26/03/01 Declarations of interest. None.

26/03/02 Minutes of 11 February approved.

26/03/03 Public participation. Mr T. Holloway raised concerns about the
condition of the footpath between the village hall and the recreation
ground. The Clerk will refer the matter to the County Council Highways
team.

26/03/04 Planning. The Council resolved to support application
24/01234/FUL (single-storey rear extension at The Old Forge) subject to
standard conditions, and to make no objection to 24/01987/LBC (sash window
replacement at 7 Church Walk).

26/03/05 Finance. Bank balance at 28 February: GBP 18,442.16. Authorised
payment of Dunbar Grounds DG-2026-0341 (GBP 561.00) and Anglia Water
26/2255198 (GBP 71.24).

26/03/06 Recreation ground play equipment - resolved to accept the
quotation from Sutcliffe Play (GBP 8,420) subject to grant funding from
WREN being confirmed.

26/03/07 Date of next meeting: Wednesday 8 April 2026, 19:30, Village Hall.""",
    )


# ---------- Agendas (3) ----------

def write_agendas():
    write_pdf(
        "agenda_full_council_2026-04-08.pdf",
        "Agenda - Full Council Meeting, 8 April 2026",
        """Meadowbrook Parish Council
Notice of meeting and agenda
Wednesday 8 April 2026, 19:30, Meadowbrook Village Hall

1. Apologies for absence
2. Declarations of interest
3. Public participation (15 minutes)
4. Minutes of the meeting held 11 March 2026
5. Planning applications - report from Clerk on current applications:
   24/02014/TPO, 24/02055/FUL, 24/02091/HOU
6. Finance:
   a. Bank reconciliation to 31 March 2026
   b. Payment of accounts
   c. Receipt of internal auditor's interim report
7. Cemetery extension - report from working group
8. Recreation ground play equipment - WREN grant outcome
9. Highways - update on School Lane verge parking
10. Annual Parish Meeting - planning for 18 March (NB date already passed - retrospective)
11. Date of next meeting: Wednesday 13 May 2026, 19:30

C Wright, Clerk to the Council
01 April 2026""",
    )
    write_pdf(
        "agenda_planning_committee_2026-04-22.pdf",
        "Agenda - Planning Committee, 22 April 2026",
        """Meadowbrook Parish Council Planning Committee
Wednesday 22 April 2026, 19:00, Meadowbrook Village Hall (small room)

1. Apologies
2. Declarations of interest
3. Minutes of meeting held 25 March 2026
4. Applications for consideration:
   a. 24/02091/HOU - 12 Willow Crescent, garage and home office
   b. 24/02112/CLD - Meadowbrook Service Station, certificate of lawfulness
   c. 24/02134/ADV - Village Hall, illuminated noticeboard
   d. 24/02158/LBC - St Mary's Church, internal re-ordering
   e. 24/02181/DOC - Primary School, discharge of conditions
   f. 24/02199/HOU - Yew Tree House, orangery
5. Recent decisions notified by the Local Planning Authority
6. Date of next meeting: Wednesday 27 May 2026, 19:00

C Wright, Clerk""",
    )
    write_pdf(
        "agenda_full_council_2026-05-13.pdf",
        "Agenda - Full Council Meeting, 13 May 2026",
        """Meadowbrook Parish Council
Annual Meeting of the Council
Wednesday 13 May 2026, 19:30, Meadowbrook Village Hall

1. Election of Chair for the municipal year 2026/27
2. Acceptance of office by the Chair
3. Election of Vice-Chair
4. Apologies for absence
5. Declarations of interest
6. Confirmation of councillors' acceptance of office and Code of Conduct
7. Appointment of committees and working groups for the year
8. Review of standing orders, financial regulations and code of conduct
9. Review of subscriptions (NALC, CAPALC)
10. Approval of insurance renewal (Zurich Municipal - GBP 1,147.44)
11. Approval of bank signatories
12. Approval of Annual Governance Statement and Annual Accounting Statement
13. Public participation
14. Minutes of meeting held 8 April 2026
15. Planning, finance, correspondence
16. Date of next meeting: Wednesday 10 June 2026

C Wright, Clerk to the Council
06 May 2026""",
    )


# ---------- Correspondence emails (5) ----------

def write_emails():
    write_eml(
        "highways_query.eml",
        "Re: Footpath condition between Village Hall and Recreation Ground",
        "Highways East <highways.east@countycouncil.gov.uk>",
        """Dear Mrs Wright,

Thank you for your enquiry of 14 March regarding the footpath between
Meadowbrook Village Hall and the recreation ground. I have logged this on
our Highways case management system as reference HW-2026-114-882.

An inspector will visit the site within the next ten working days and we
will write to you again with our findings and proposed remedial action.

In the meantime please advise residents to take care, particularly in wet
conditions where the surface degradation is most pronounced.

Kind regards,
James Allcock
Highways East - Footways Team""",
        when="2026-03-18",
    )
    write_eml(
        "wren_grant_decision.eml",
        "WREN Community Action Fund - Grant Decision MB-PC-2026-0184",
        "WREN Grant Officer <grants@wren.org.uk>",
        """Dear Mrs Wright,

I am writing to confirm that the trustees of WREN's Community Action
Fund have approved your application for funding towards the replacement
of play equipment at Meadowbrook Recreation Ground.

Reference:    MB-PC-2026-0184
Award:        GBP 7,500 (towards a project total of GBP 8,420)
Conditions:   Works to be completed by 31 December 2026; standard WREN
              monitoring and reporting requirements apply; the WREN logo
              and acknowledgement must be displayed on the equipment.

Please countersign the offer letter (attached to the original email) and
return it within 28 days to formally accept the grant.

Best regards,
Philippa Stenton
Grants Officer, WREN Community Action Fund""",
        when="2026-04-02",
    )
    write_eml(
        "supplier_query_kompan.eml",
        "Quotation for replacement swing frame - reference QU/MB/2026/0044",
        "Sales (Kompan UK) <sales-uk@kompan.com>",
        """Dear Mrs Wright,

Following our site visit on 18 February please find below the quotation
for the supply and installation of one Kompan multi-swing frame at
Meadowbrook Recreation Ground.

Item: M204 Robinia multi-swing frame, two flat seats and one cradle.
Supply, delivery, removal of existing equipment, installation, surfacing
and post-installation report.

Total ex VAT:    GBP 8,955
VAT (20%):       GBP 1,791
TOTAL inc VAT:   GBP 10,746

Quotation valid until 30 April 2026. Lead time 8-10 weeks from order.

Kind regards,
Caroline Etheridge
Senior Account Manager, Kompan UK""",
        when="2026-02-24",
    )
    write_eml(
        "audit_engagement_letter.eml",
        "Internal audit engagement letter 2025/26",
        "Goodwin & Hayes <audit@goodwinhayes.co.uk>",
        """Dear Mrs Wright,

Please find attached our engagement letter for the internal audit of
Meadowbrook Parish Council for the year ended 31 March 2026.

Our work will follow the Practitioners' Guide for Local Councils 2024
and will include verification of the items required for completion of
the Annual Internal Audit Report on the AGAR.

Proposed fieldwork: w/c 14 April 2026 (one day on site).
Fee: GBP 480 plus disbursements.

Please countersign the engagement letter and return one copy.

Yours sincerely,
Daniel Hayes FCCA
Goodwin & Hayes Audit Services""",
        when="2026-03-26",
    )
    write_eml(
        "ico_subject_access.eml",
        "Subject access request - acknowledgement",
        "Information Commissioner's Office <casework@ico.org.uk>",
        """Dear Mrs Wright,

We acknowledge receipt of the subject access request you forwarded on
22 March 2026 from a former resident regarding correspondence held by
the Council in 2019.

The Council has 30 calendar days from the date the request was first
received (15 March 2026) to respond. We have logged this case as
IC-2026-MB-0008812 for your reference.

If you require any guidance on responding, please contact our helpline
or use the SAR self-assessment tool on our website.

Regards,
Information Commissioner's Office Casework Team""",
        when="2026-03-23",
    )


# ---------- Reports (2) ----------

def write_reports():
    write_pdf(
        "annual_return_AGAR_2024-25.pdf",
        "Annual Governance and Accountability Return 2024/25 - Section 1 and Section 2",
        """Meadowbrook Parish Council
Smaller Authority Annual Governance and Accountability Return
For the year ended 31 March 2025

Section 1 - Annual Governance Statement
The Council confirms the following with respect to the year ended
31 March 2025:

1. We have put in place arrangements for effective financial management
   during the year, and for the preparation of the accounting statements.
   Yes.
2. We have maintained an adequate system of internal control. Yes.
3. We have taken all reasonable steps to assure ourselves that there are
   no matters of actual or potential non-compliance with laws,
   regulations and proper practices. Yes.
4. We provided proper opportunity for the exercise of electors' rights
   in accordance with the requirements of the Accounts and Audit
   Regulations. Yes.
5. We carried out an assessment of the risks facing the Council. Yes.
6. We maintained throughout the year an adequate and effective system of
   internal audit. Yes.
7. We took appropriate action on all matters raised in reports from
   internal and external audit. Yes.
8. We considered whether any litigation, liabilities or commitments,
   events or transactions occurring either during or after the year-end
   have a financial impact on the Council. None.
9. The Council has met all its responsibilities where it is a sole
   managing trustee of a local trust or trusts. Not applicable.

Section 2 - Accounting Statements (GBP)

                                            2023/24    2024/25
1. Balances brought forward                  17,205     19,210
2. (+) Precept or rates and levies           38,500     41,200
3. (+) Total other receipts                   3,920      4,485
4. (-) Staff costs                            8,420      8,650
5. (-) Loan interest / capital repayments         0          0
6. (-) Total other payments                  31,995     38,803
7. (=) Balances carried forward              19,210     17,442
8. Total value of cash and short-term inv    19,210     17,442
9. Total fixed assets and long-term inv      88,400     88,920
10. Total borrowings                              0          0

Signed: S Whitfield (Chair) and C Wright (Clerk/RFO)
Date: 6 May 2025""",
    )
    write_pdf(
        "asset_register_2025-03-31.pdf",
        "Asset Register at 31 March 2025",
        """Meadowbrook Parish Council
Asset Register at 31 March 2025

Land and buildings
- Village Green (registered title CB117422) - nominal GBP 1
- Cemetery (registered title CB214991) - nominal GBP 1
- Recreation Ground (registered title CB551200) - nominal GBP 1
- Pavilion building, Recreation Ground - GBP 42,000 (insured value)

Equipment and street furniture
- Memorial benches x 4 - GBP 2,400
- Litter bins x 11 - GBP 880
- Notice boards x 3 - GBP 1,650
- Bus shelter, High Street - GBP 3,200
- Christmas tree lighting set - GBP 480
- Defibrillator (cabinet at Pavilion) - GBP 1,250

Recreation equipment
- Multi-swing frame (1986, due replacement 2026) - GBP 0 NBV
- Steel slide and platform - GBP 1,800
- Cricket sightscreens (existing) - GBP 950
- Cricket pitch covers - GBP 720

Office equipment
- Clerk's laptop (Lenovo, 2023) - GBP 480
- Printer / scanner - GBP 220
- Filing cabinets x 2 - GBP 180

Total recorded asset value (insurance-equivalent):  GBP 56,213

Notes
- Land titles held at nominal GBP 1 in line with NALC guidance
- Insurance valuations reviewed annually with Zurich Municipal
- Asset register reviewed and approved at Annual Meeting""",
    )


def main():
    write_planning_notices()
    write_invoices()
    write_minutes()
    write_agendas()
    write_emails()
    write_reports()
    print(f"[sample-docs-gen] wrote {len(os.listdir(OUT_DIR))} files to {OUT_DIR}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
